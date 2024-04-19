from django.http import HttpResponse,FileResponse
from utils import device_coding
import json
from docx import Document
from UserManagement.role_service import findPermissionByName,findRoleTypeByName
from py2neo import Node,Relationship,Graph,Path,Subgraph
from py2neo import NodeMatcher,RelationshipMatcher
from datetime import datetime,timedelta
from time import sleep
import time
import re

# 操作excel文件
from openpyxl import Workbook
from openpyxl.styles import Alignment
from openpyxl.utils import get_column_letter
from io import BytesIO  




neo4j_url='bolt://211.82.97.250:7688'
user='neo4j'
password='neo4j'
graph=Graph(neo4j_url,auth=(user,password))



# 查询有编码设备情况
def queryDeviceNumber(request):
    
    def get_num(num_str):
        sum_ = 0
        try:
            number = re.findall("\d+", num_str)  # 输出结果为列表
            for num_ in number:
                sum_ += int(num_)
        except Exception:
            sum_=None
        finally:
            return sum_
        

    if request.method=='POST':
        # postbody=request.body
        # param=json.loads(postbody.decode())
        # name=param['name']
        try:
            sql = "MATCH (n) WHERE (n.UniqueCode) IS NOT NULL RETURN DISTINCT n.UniqueCode AS UniqueCode,n.数量 as number; "
            datas = graph.run(sql).data()   # 所有设备编码组成的列表 [{'UniqueCode': '0109090001'}, {'UniqueCode': '0109040001'}, {'UniqueCode': '0109040002'},
            types = {
                '01' : '采煤工作面设备',
                '02' : '掘进工作面设备',
                '04' : '辅助运输设备',
                '05' : '通风与压风设备',
                '07' : '安全监控设备',
                '03' : '主煤流运输设备',
                '08' : '绿色、节能设备',
                '06' : '供电与给排水设备',
                '09' : '其他系统机电设备'}

            companys = {
                '02' : '扎赉诺尔煤业公司',
                '04' : '北方公司',
                '05' : '庆阳煤电公司',
                '03' : '陕西矿业分公司',
                '01' : '华亭煤业公司',
                '06' : '滇东矿业分公司'}
            
            # 初始化存储数据格式
            devices = {}
            for company in companys.values():
                devices[company]={}
                devices[company]['total'] = 0
                for type in types.values():
                    devices[company][type] = 0
            
            for data in datas:
                uniquecode = data['UniqueCode']
                num = get_num(data['number'])
                company_key = uniquecode[0:2]
                type_key = uniquecode[4:6]
                devices[companys[company_key]][types[type_key]]+=num
                devices[companys[company_key]]['total']+=num
        

            res={
                'code' : 10000,
                'equipment_company_num' : devices
            }
        except Exception:
            res={
                'code':10001,
                'equipment_company_num' : None
            }
        
        finally:

            return HttpResponse(json.dumps(res,ensure_ascii=False))

# 查询有编码设备情况V2，并返回前端统计展示
def queryDeviceNumber2(request):
    
    def get_num(num_str):
        sum_ = 0
        try:
            number = re.findall("\d+", num_str)  # 输出结果为列表
            for num_ in number:
                sum_ += int(num_)
        except Exception:
            sum_=None
        finally:
            return sum_
        
    def devices_num_for_coal(coals,types,datas,role_name):
        devices = {}
        devices['total'] = 0
        for type in types.values():
            devices[type] = 0
        for data in datas:  # 遍历所有设备
            uniquecode = data['UniqueCode']
            num = get_num(data['number'])  # 一个设备的数量
            coal_key = uniquecode[2:4]  # 所属矿山
            type_key = uniquecode[4:6]  # 所属工作面
        # print(coals[coal_key])
            if type_key=='10':
                type_key='09'   # 把洗选煤归属到 其他
            if type_key in types.keys() and type_key in types.keys():
                if coals[coal_key] == role_name:
                    devices[types[type_key]] += num
                    devices['total'] += num
                
        # print(devices)
        return {role_name: devices}  # 满足前端
    def devices_num_for_headOffice(companys,types,datas):
        """
        查找总公司设备数量
        :param companys:
        :param types:
        :param datas:
        :return:
        """
        devices = {}
        for company in companys.values():
            devices[company] = {}
            devices[company]['total'] = 0
            for type in types.values():
                devices[company][type] = 0
        
        for data in datas: # 遍历所有设备
            uniquecode = data['UniqueCode']
            num = get_num(data['number'])  # 一个设备的数量
            company_key = uniquecode[0:2]  # 所属公司
            type_key = uniquecode[4:6]     # 所属工作面
            if type_key=='10':
                type_key='09'   # 把洗选煤归属到 其他
            if type_key in types.keys() and company_key in companys.keys():
                devices[companys[company_key]][types[type_key]] += num
                devices[companys[company_key]]['total'] += num
        return devices
    
    def devices_num_for_company(coals,types,datas,role_name):
        """
        查找分公司设备数量
        """
        sub_coals=findPermissionByName(role_name=role_name)
        devices = {}
        for coal in sub_coals:
            devices[coal] = {}
            devices[coal]['total'] = 0
            for type in types.values():
                devices[coal][type] = 0
        
        for data in datas:  # 遍历所有设备

            uniquecode = data['UniqueCode']
            num = get_num(data['number'])  # 一个设备的数量
            coal_key = uniquecode[2:4]  # 所属矿山
            type_key = uniquecode[4:6]  # 所属工作面
            if type_key=='10':
                type_key='09'   # 把洗选煤归属到 其他
            if type_key in types.keys() and type_key in types.keys():
                if coals[coal_key] in sub_coals:
                    devices[coals[coal_key]][types[type_key]] += num
                    devices[coals[coal_key]]['total'] += num
        
        return {role_name:devices}  # 满足前端
    
    if request.method=='POST':
        postbody=request.body
        param = json.loads(postbody.decode(), encoding='utf-8')
        role_name=param['role']
        role_type=findRoleTypeByName(role_name=role_name)
        # types = {
        #         '01' : '采煤工作面设备',
        #         '02' : '掘进工作面设备',
        #         '04' : '辅助运输设备',
        #         '05' : '通风与压风设备',
        #         '07' : '安全监控设备',
        #         '03' : '主煤流运输设备',
        #         '08' : '绿色、节能设备',
        #         '06' : '供电与给排水设备',
        #         '09' : '其他系统机电设备'}

        types=device_coding.type_coding()
        # companys = {
        #         '02' : '扎赉诺尔煤业公司',
        #         '04' : '北方公司',
        #         '05' : '庆阳煤电公司',
        #         '03' : '陕西矿业分公司',
        #         '01' : '华亭煤业公司',
        #         '06' : '滇东矿业分公司'}
        
        companys = device_coding.company_coding()
    #     coals={
    # '01':'华亭煤矿','02':'砚北煤矿','03':'陈家沟煤矿','04':'东峡煤矿','05':'山寨煤矿',
    # '06':'马蹄沟煤矿','07':'新窑煤矿','08':'新柏煤矿','09':'大柳煤矿','10':'赤城煤矿',
    # '11':'灵东煤矿','12':'灵泉煤矿','13':'铁北煤矿','14':'灵露煤矿','15':'青岗坪煤矿',
    # '16':'柳巷煤矿','17':'西川煤矿','18':'高头窑煤矿','19':'新庄煤矿','20':'核桃峪煤矿',
    # '21':'刘园子煤矿','22':'白龙山煤矿一井','23':'雨汪煤矿一井'}
        coals = device_coding.coal_coding()

        try:
            sql = "MATCH (n) WHERE (n.UniqueCode) IS NOT NULL RETURN DISTINCT n.UniqueCode AS UniqueCode,n.数量 as number; "
            datas = graph.run(sql).data()   # 所有设备编码组成的列表 [{'UniqueCode': '0109090001'}, {'UniqueCode': '0109040001'}, {'UniqueCode': '0109040002'},
            devices={}
            if role_type==0:
                devices=devices_num_for_headOffice(companys,types,datas)

            elif role_type==1:
                devices=devices_num_for_company(coals,types,datas,role_name)
            elif role_type==2:
                devices=devices_num_for_coal(coals,types=types,datas=datas,role_name=role_name)
        
            res={
                'code' : 10000,
                'equipment_company_num' : devices
            }
        except Exception:
            res={
                'code':10001,
                'equipment_company_num' : None
            }
        
        finally:

            return HttpResponse(json.dumps(res,ensure_ascii=False))
        
# 前端表格展示
def showNodeForTable(request):
    if request.method=='POST':
        postbody=request.body
        param=json.loads(postbody.decode(),encoding='utf-8')

    # 保存前端传入数据为txt，方便查看    
        # file = open('test.txt','w')
        # ascii不能输出中文
        # str = json.dumps(param,ensure_ascii=False)
        # file.write(str)
        # file.close()
    # 。。。。。。。。。。。。。。。。。。。。。。。。。。。。
        node_list=param['node_list']
        rel_list=param['rel_list']
        # 写一个分级保存的框架，按照假节点进行分类
        # table_labels = ['类型',
        #         '名称',
        #         '开始使用日期',
        #         '出厂编码',
        #         '存放地点',
        #         '所属煤业',
        #         '所属矿山',
        #         '所属集团',
        #         '技术参数',
        #         '数量',
        #         '生产厂家',
        #         '设备号',
        #         '设备状态',
        #         '设备编码',
        #         '责任人',
        #         'UniqueCode']
        table_labels=[]
        # 数据分级
        # []->  {type1:[],type2:[],type3:[]}
        # 进行数据处理
        types = {
            '01' : '采煤工作面设备',
            '02' : '掘进工作面设备',
            '03' : '主煤流运输设备',
            '04' : '辅助运输设备',
            '05' : '通风与压风设备',
            '06' : '供电与给排水设备',
            '07' : '安全监控设备',
            '08' : '绿色、节能设备',
            '09' : '其他系统机电设备'}
        
        nodes = {}
        for node in node_list:
            data={}
            d={}
            if node.__contains__('properites') and node['properites'].get('UniqueCode') is not None and len(node['properites'].get('UniqueCode')) == 10:
                if len(table_labels)==0: # 动态生成表头
                    table_labels.append('设备类型')
                    table_labels.append('名称')
                    for key in node['properites'].keys():
                        table_labels.append(key)
                    
                # 只存有编码的设备
                uniquecode = node['properites'].get('UniqueCode')
                type_key = uniquecode[4:6]
                data['设备类型'] = types.get(type_key)  # 存入所属类型
                data['名称'] = node['label']
                d = {**data, **node['properites']}

                if nodes.__contains__(d['设备类型']):
                    # 对节点进行分类别存储
                    nodes[d['设备类型']].append(d)
                else:
                    nodes[d['设备类型']] = [d]
            # 数据处理完成 []->{type1:[ [],[] ], type2:[[],[]]...}
        
        # 格式转换
        out=[]
        for node in nodes.values():
            out = out+node
        res={
            'table_labels' : table_labels,
            'node' : out,
        }
        # print("表格展示数据：：")
        # print(res)
        return HttpResponse(json.dumps(res,ensure_ascii=False))

#保存节点成csv文件
def saveNodeToCSV(request):
    if request.method=='POST':
        postbody = request.body
        param = json.loads(postbody.decode(),encoding='utf-8')
        
    # 保存前端传入数据为txt，方便查看    
        file = open('test.txt','w')
        # ascii不能输出中文
        str = json.dumps(param,ensure_ascii=False)
        file.write(str)
        file.close()
    # 。。。。。。。。。。。。。。。。。。。。。。。。。。。。
        node_list = param['node_list']
        rel_list = param['rel_list']
        # 写一个分级保存的框架，按照假节点进行分类
        wb = Workbook()
        ws = wb.active
        # table_labels = ['类型',
        #         '名称',
        #         '开始使用日期',
        #         '出厂编码',
        #         '存放地点',
        #         '所属煤业',
        #         '所属矿山',
        #         '所属集团',
        #         '技术参数',
        #         '数量',
        #         '生产厂家',
        #         '设备号',
        #         '设备状态',
        #         '设备编码',
        #         '责任人',
        #         'UniqueCode']
        table_labels=[]
        # 数据分级
        # []->  {type1:[],type2:[],type3:[]}
        # ws.append(table_labels)
        # 进行数据处理
        types = {
            '01' : '采煤工作面设备',
            '02' : '掘进工作面设备',
            '03' : '主煤流运输设备',
            '04' : '辅助运输设备',
            '05' : '通风与压风设备',
            '06' : '供电与给排水设备',
            '07' : '安全监控设备',
            '08' : '绿色、节能设备',
            '09' : '其他系统机电设备'}
        
        nodes = {}
        for node in node_list:
            data=[]
            if node.__contains__('properites') and node['properites'].get('UniqueCode') is not None and len(node['properites'].get('UniqueCode')) == 10:
                if len(table_labels)==0:
                    table_labels.append('设备类型')
                    table_labels.append('名称')
                    for key in node['properites'].keys():
                        table_labels.append(key)
                    ws.append(table_labels)
                # 只存有编码的设备
                uniquecode = node['properites'].get('UniqueCode')
                type_key = uniquecode[4:6]
                data.append(types.get(type_key))  # 存入所属类型
                data.append(node['label'])
                for value in node['properites'].values():
                    data.append(value)
                if nodes.__contains__(data[0]):
                    # 对节点进行分类别存储
                    nodes[data[0]].append(data)
                else:
                    nodes[data[0]]=[data]
            # 数据处理完成 []->{type1:[ [],[] ], type2:[[],[]]...}
                        
        for key in nodes.keys():
            max_rows1 = ws.max_row  # 获取本类插入前最大行
            for data in nodes.get(key): # 遍历一类
                ws.append(data)
            max_rows2 = ws.max_row  # 获取本类插入后最大行
            # 合并
            ws.merge_cells('A{}:A{}'.format(max_rows1+1,max_rows2))
        
        # 居中设置
        max_rows = ws.max_row  # 获取最大行
        max_columns = ws.max_column  # 获取最大列
        align=Alignment(horizontal='center',vertical='center')
        # openpyxl的下标从1开始
        for i in range(1, max_columns + 1):
            if i<3:
                ws.column_dimensions[get_column_letter(i)].width = 40
            else:
                ws.column_dimensions[get_column_letter(i)].width = 10
            for j in range(1, max_rows + 1):
                ws.cell(j, i).alignment = align
        t = time.localtime()
        file_name = "{}{}{}{}{}{}.xlsx".format(t.tm_year,t.tm_mon,t.tm_mday,t.tm_hour,t.tm_min,t.tm_sec)
        # sio = BytesIO()
        wb.save("temp.xlsx")
        # tooken={
        #     'code' : 200,
        # }
        # return HttpResponse(json.dumps(tooken,ensure_ascii=False))
        # wb.save(sio)
        # sio.seek(0)
        file = open('temp.xlsx', 'rb')
        response = FileResponse(file,as_attachment=True)
        response['Content-Type'] = 'application/octet-stream'
        response['Content-Disposition'] = 'attachment;filename="test.xlsx"'
        
        # response.write(sio.getvalue())
        return response

# 保存节点成world文件
def saveNodeToWorld(request): 
    if request.method == 'POST':
        postbody = request.body
        param = json.loads(postbody.decode(), encoding='utf-8')

        # 保存前端传入数据为txt，方便查看
        file = open('test.txt', 'w')
        # ascii不能输出中文
        str = json.dumps(param, ensure_ascii=False)
        file.write(str)
        file.close()
        # 。。。。。。。。。。。。。。。。。。。。。。。。。。。。
        node_list = param['node_list']
        rel_list = param['rel_list']
        # 写一个分级保存的框架，按照假节点进行分类
        # wb = Workbook()
        # ws = wb.active
        # table_labels = ['类型',
        #         '名称',
        #         '开始使用日期',
        #         '出厂编码',
        #         '存放地点',
        #         '所属煤业',
        #         '所属矿山',
        #         '所属集团',
        #         '技术参数',
        #         '数量',
        #         '生产厂家',
        #         '设备号',
        #         '设备状态',
        #         '设备编码',
        #         '责任人',
        #         'UniqueCode']
        table_labels = []
        # 数据分级
        # []->  {type1:[],type2:[],type3:[]}
        # ws.append(table_labels)
        # 进行数据处理
        types = {
            '01': '采煤工作面设备',
            '02': '掘进工作面设备',
            '03': '主煤流运输设备',
            '04': '辅助运输设备',
            '05': '通风与压风设备',
            '06': '供电与给排水设备',
            '07': '安全监控设备',
            '08': '绿色、节能设备',
            '09': '其他系统机电设备'}

        nodes = {}

        for node in node_list:
            data = []
            if node.__contains__('properites') and node['properites'].get('UniqueCode') is not None and len(
                    node['properites'].get('UniqueCode')) == 10:
                if len(table_labels) == 0:
                    table_labels.append('设备类型')
                    table_labels.append('名称')
                    for key in node['properites'].keys():
                        table_labels.append(key)
                # -----
                # 只存有编码的设备
                uniquecode = node['properites'].get('UniqueCode')
                type_key = uniquecode[4:6]
                data.append(types.get(type_key))  # 存入所属类型
                data.append(node['label'])
                for value in node['properites'].values():
                    data.append(value)
                if nodes.__contains__(data[0]):
                    # 对节点进行分类别存储
                    nodes[data[0]].append(data)
                else:
                    nodes[data[0]] = [data]
            # 数据处理完成 []->{type1:[ [],[] ], type2:[[],[]]...}

        # for key in nodes.keys():
        #     max_rows1 = ws.max_row  # 获取本类插入前最大行
        #     for data in nodes.get(key):  # 遍历一类
        #         ws.append(data)
        #     max_rows2 = ws.max_row  # 获取本类插入后最大行
        #     # 合并
        #     ws.merge_cells('A{}:A{}'.format(max_rows1 + 1, max_rows2))

        doc = Document()  # 创建一个Document对象

        #-------- 测试使用--------------------------
        # doc_test = Document()
        # doc_test.add_paragraph('这是一个段落。')
        # doc_test.save("temp_test.docx")

        # doc_table_test = Document()
        # doc_table_test.add_table(rows=1, cols=len(table_labels), style='Table Grid')
        # doc_table_test.save("temp_table_test.docx")
        # ---------------
        if len(table_labels)>0:
            table = doc.add_table(rows=1, cols=len(table_labels), style='Table Grid')
            table.autofit = True
            row0 = table.rows[0]
            for i in range(0, len(row0.cells)):
                row0.cells[i].text = table_labels[i]  # 插入table label

            for key in nodes.keys():
                max_rows1 = len(table.rows)  # 获取本类插入前最大行
                for data in nodes.get(key):  # 遍历一类
                    row_cells = table.add_row().cells
                    for i in range(0, len(row_cells)):
                        if data[i] is None:
                            data[i] = ''
                        row_cells[i].text = data[i]  # 插入table label
                max_rows2 = len(table.rows)  # 获取本类插入后最大行

            # 居中设置
            # max_rows = ws.max_row  # 获取最大行
            # max_columns = ws.max_column  # 获取最大列
            # align = Alignment(horizontal='center', vertical='center')
            # # openpyxl的下标从1开始
            # for i in range(1, max_columns + 1):
            #     if i < 3:
            #         ws.column_dimensions[get_column_letter(i)].width = 40
            #     else:
            #         ws.column_dimensions[get_column_letter(i)].width = 10
            #     for j in range(1, max_rows + 1):
            #         ws.cell(j, i).alignment = align
            # t = time.localtime()
            # file_name = "{}{}{}{}{}{}.xlsx".format(t.tm_year, t.tm_mon, t.tm_mday, t.tm_hour, t.tm_min, t.tm_sec)
            # sio = BytesIO()
            # wb.save("temp.xlsx")
        doc.save("temp.docx")

        # tooken={
        #     'code' : 200,
        # }
        # return HttpResponse(json.dumps(tooken,ensure_ascii=False))
        # wb.save(sio)
        # sio.seek(0)
        file = open('temp.docx', 'rb')
        response = FileResponse(file, as_attachment=True)
        # response['Content-Type'] = 'application/octet-stream'
        # response['Content-Type'] = 'application/msword'
        response['Content_Type']='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response['Content-Disposition'] = 'attachment;filename="temp.docx"'

        # response.write(sio.getvalue())
        return response